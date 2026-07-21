from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass, field
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any

from boardsight_ai.models import PipelineResult


BRAND_NAVY = "#0F172A"
BRAND_BLUE = "#2563EB"
BRAND_INDIGO = "#4F46E5"
BRAND_GREEN = "#16A34A"
BRAND_AMBER = "#D97706"
BRAND_RED = "#DC2626"
BRAND_TEXT = "#1E293B"
BRAND_MUTED = "#64748B"
BRAND_BORDER = "#CBD5E1"
BRAND_SURFACE = "#F8FAFC"

WEEKDAY_INDEX = {
    "monday": 0,
    "tuesday": 1,
    "wednesday": 2,
    "thursday": 3,
    "friday": 4,
    "saturday": 5,
    "sunday": 6,
}


@dataclass
class ReportDecisionRow:
    decision_id: str
    title: str
    exact_text: str
    owner: str
    timestamp: str
    rationale: str
    evidence: str
    urgency: str
    impact: str
    blockers: str
    next_action: str
    status: str
    linked_gitlab_issue: str
    priority_score: float
    confidence: float


@dataclass
class ReportActionRow:
    action_id: str
    task_title: str
    derived_from_decision_id: str
    owner: str
    due_date: str
    confidence: float
    dependencies: str
    blocker_flag: str
    gitlab_sync_status: str
    linked_gitlab_issue: str
    status: str
    execution_order: int
    notes: str


@dataclass
class ReportRiskRow:
    risk_id: str
    category: str
    severity: str
    description: str
    owner: str
    related_decision_id: str
    recommended_follow_up: str


@dataclass
class ReportWorkflowRow:
    stage: str
    outbound_transition: str
    stage_note: str


@dataclass
class ReportSpeakerRow:
    speaker: str
    dominance_ratio: float
    rating: str
    participation_note: str


@dataclass
class ReportVisualRow:
    artifact_id: str
    artifact_type: str
    display_mode: str
    time_window: str
    summary: str
    insight: str
    confidence: float


@dataclass
class ReportTraceRow:
    trace_id: str
    title: str
    owner: str
    summary: str
    rationale: str
    next_steps: str
    related_artifacts: str


@dataclass
class StructuredReportModel:
    report_title: str
    generated_at: str
    executive_summary: list[str]
    metadata_rows: list[tuple[str, str]]
    decisions: list[ReportDecisionRow]
    actions: list[ReportActionRow]
    risks: list[ReportRiskRow]
    workflow: list[ReportWorkflowRow]
    participation: list[ReportSpeakerRow]
    visuals: list[ReportVisualRow]
    traces: list[ReportTraceRow]
    recommended_follow_through: list[str]
    warnings: list[str] = field(default_factory=list)


def _coalesce_mapping_value(mapping: dict[str, object], *keys: str, default: object = "n/a") -> object:
    for key in keys:
        if key in mapping and mapping[key] not in (None, ""):
            return mapping[key]
    return default


def _safe_text(value: Any, default: str = "n/a") -> str:
    text = str(value or "").strip()
    return text or default


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _status_badge(status: str) -> str:
    return status.replace("-", " ").title()


def _compact_text(value: str, limit: int = 96) -> str:
    text = " ".join(str(value or "").split())
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def _priority_band(score: float) -> tuple[str, str]:
    if score >= 0.85:
        return "High", "High"
    if score >= 0.6:
        return "Medium", "Medium"
    return "Low", "Low"


def _normalize_due_date(text: str) -> str:
    lowered = text.lower()
    for name, weekday in WEEKDAY_INDEX.items():
        if name in lowered:
            now = datetime.now(UTC)
            days_ahead = (weekday - now.weekday()) % 7
            days_ahead = 7 if days_ahead == 0 else days_ahead
            return (now + timedelta(days=days_ahead)).date().isoformat()
    match = re.search(r"\b(\d{4}-\d{2}-\d{2})\b", text)
    if match:
        return match.group(1)
    return ""


def _extract_gitlab_link(entry: dict[str, Any]) -> str:
    for key in ("gitlab_issue_url", "issue_web_url", "web_url", "issue_url"):
        value = str(entry.get(key) or "").strip()
        if value:
            return value
    if entry.get("issue_iid"):
        return f"Issue #{entry['issue_iid']}"
    return ""


def _render_list(items: list[Any], default: str = "None") -> str:
    values = [str(item).strip() for item in items if str(item).strip()]
    return "; ".join(values) if values else default


def _format_timestamp_range(start: Any, end: Any) -> str:
    return f"{_safe_float(start):.1f}s - {_safe_float(end):.1f}s"


def _derive_title_from_text(text: str, fallback: str) -> str:
    compact = _compact_text(text, limit=68)
    return compact if compact and compact != "n/a" else fallback


def _format_ratio(value: float) -> str:
    return f"{value:.1f}%"


def _infer_action_status(owner: str, due_date: str, blocker_flag: str, linked_issue: str) -> str:
    if blocker_flag == "Yes":
        return "Blocked"
    if owner == "Unassigned":
        return "Needs owner"
    if not due_date:
        return "Needs due date"
    if linked_issue:
        return "Synced"
    return "Ready"


def _add_missing_follow_through_risks(
    actions: list[ReportActionRow],
    decisions: list[ReportDecisionRow],
    risks: list[ReportRiskRow],
) -> None:
    existing_ids = {item.risk_id for item in risks}
    if any(action.owner == "Unassigned" for action in actions) and "RISK-OWNER" not in existing_ids:
        risks.append(
            ReportRiskRow(
                risk_id="RISK-OWNER",
                category="Missing owner",
                severity="High",
                description="One or more action items do not yet have a clear owner.",
                owner="PMO",
                related_decision_id="",
                recommended_follow_up="Assign named owners before closing the meeting.",
            )
        )
    if any(not action.due_date for action in actions) and "RISK-DATE" not in existing_ids:
        risks.append(
            ReportRiskRow(
                risk_id="RISK-DATE",
                category="Missing deadline",
                severity="Medium",
                description="One or more action items are missing a usable due date.",
                owner="PMO",
                related_decision_id="",
                recommended_follow_up="Capture due dates for each open action item.",
            )
        )
    if any(decision.status in {"Captured", "Needs owner"} for decision in decisions) and "RISK-UNRESOLVED" not in existing_ids:
        risks.append(
            ReportRiskRow(
                risk_id="RISK-UNRESOLVED",
                category="Unresolved decision",
                severity="Medium",
                description="At least one decision was captured without complete follow-through metadata.",
                owner="Meeting owner",
                related_decision_id="",
                recommended_follow_up="Confirm next step, owner, and execution status for open decisions.",
            )
        )


def build_structured_report_model(result: PipelineResult) -> StructuredReportModel:
    generated_at = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    metadata = result.metadata or {}
    workflow = result.workflow_model
    agentic_contract = metadata.get("agentic_contract") if isinstance(metadata.get("agentic_contract"), dict) else {}
    contract_entities = agentic_contract.get("entities") if isinstance(agentic_contract.get("entities"), dict) else {}
    risk_signals = contract_entities.get("risk_signals") if isinstance(contract_entities.get("risk_signals"), list) else []

    decision_trace_lookup = {
        item.trace_id: item for item in result.decision_traces
    }
    priority_lookup = {
        str(item.get("decision_id") or ""): item for item in workflow.prioritized_decisions if str(item.get("decision_id") or "").strip()
    }
    tasks_by_decision: dict[str, list[dict[str, Any]]] = {}
    for task in workflow.execution_plan:
        decision_id = str(task.get("decision_id") or "").strip()
        if decision_id:
            tasks_by_decision.setdefault(decision_id, []).append(task)

    trace_by_decision: dict[str, Any] = {}
    for trace in result.decision_traces:
        if trace.execution_tasks:
            for task in trace.execution_tasks:
                decision_id = str(task.get("decision_id") or "").strip()
                if decision_id:
                    trace_by_decision.setdefault(decision_id, trace)
        if trace.trace_id and trace.trace_id.startswith("TRACE-"):
            fallback_id = trace.trace_id.replace("TRACE", "DM", 1)
            trace_by_decision.setdefault(fallback_id, trace)

    decisions: list[ReportDecisionRow] = []
    for index, moment in enumerate(result.decision_moments, start=1):
        decision_id = moment.event_id or f"DEC-{index}"
        priority_entry = priority_lookup.get(decision_id, {})
        trace = trace_by_decision.get(decision_id)
        if trace is None:
            trace = next(
                (
                    item
                    for item in result.decision_traces
                    if decision_id in " ".join(item.next_steps + item.rationale + [item.summary, item.title])
                ),
                None,
            )
        linked_tasks = list(tasks_by_decision.get(decision_id, []))
        priority_score = _safe_float(priority_entry.get("priority_score"), default=moment.confidence)
        urgency, impact = _priority_band(priority_score)
        blockers = [
            text
            for text in workflow.bottlenecks
            if any(keyword in text.lower() for keyword in decision_id.lower().split("-") + moment.text.lower().split()[:3])
        ]
        if not blockers:
            blockers = list(workflow.bottlenecks[:1])
        next_action = linked_tasks[0].get("title") if linked_tasks else (trace.next_steps[0] if trace and trace.next_steps else "")
        owner = _safe_text(
            (trace.owner if trace and trace.owner else "") or linked_tasks[0].get("owner") if linked_tasks else moment.speaker,
            default="Unassigned",
        )
        linked_issue = _extract_gitlab_link(linked_tasks[0]) if linked_tasks else ""
        status = "Captured"
        if blockers:
            status = "Blocked"
        elif owner == "Unassigned":
            status = "Needs owner"
        elif next_action:
            status = "Ready"
        decisions.append(
            ReportDecisionRow(
                decision_id=decision_id,
                title=_derive_title_from_text(trace.title if trace and trace.title else moment.text, f"Decision {index}"),
                exact_text=_safe_text(moment.text),
                owner=owner,
                timestamp=_safe_text(moment.timestamp, default=f"{index}"),
                rationale=_render_list(trace.rationale if trace else []),
                evidence=_render_list(moment.evidence),
                urgency=urgency,
                impact=impact,
                blockers=_render_list(blockers),
                next_action=_safe_text(next_action, default="No immediate follow-through recorded"),
                status=_status_badge(status),
                linked_gitlab_issue=linked_issue or "Not synced",
                priority_score=priority_score,
                confidence=_safe_float(moment.confidence),
            )
        )

    action_rows: list[ReportActionRow] = []
    for index, task in enumerate(workflow.execution_plan, start=1):
        linked_issue = _extract_gitlab_link(task)
        notes = _safe_text(task.get("notes") or task.get("text") or "", default="")
        due_date = _safe_text(task.get("due_date") or _normalize_due_date(f"{task.get('title', '')} {notes}"), default="")
        blocker_flag = "Yes" if ("depend" in notes.lower() or "block" in notes.lower()) else "No"
        action_rows.append(
            ReportActionRow(
                action_id=_safe_text(task.get("task_id"), default=f"ACTION-{index}"),
                task_title=_safe_text(task.get("title"), default=f"Action item {index}"),
                derived_from_decision_id=_safe_text(task.get("decision_id"), default=""),
                owner=_safe_text(task.get("owner"), default="Unassigned"),
                due_date=due_date,
                confidence=_safe_float(task.get("priority_score"), default=0.0),
                dependencies=_safe_text(notes if blocker_flag == "Yes" else "", default="None"),
                blocker_flag=blocker_flag,
                gitlab_sync_status="Synced" if linked_issue else "Not synced",
                linked_gitlab_issue=linked_issue or "",
                status=_infer_action_status(
                    _safe_text(task.get("owner"), default="Unassigned"),
                    due_date,
                    blocker_flag,
                    linked_issue,
                ),
                execution_order=int(task.get("execution_order") or index),
                notes=notes or "No implementation note recorded.",
            )
        )

    risks: list[ReportRiskRow] = []
    for index, item in enumerate(workflow.bottlenecks, start=1):
        risks.append(
            ReportRiskRow(
                risk_id=f"BLOCKER-{index}",
                category="Blocker",
                severity="High",
                description=_safe_text(item),
                owner="Meeting owner",
                related_decision_id="",
                recommended_follow_up="Resolve the blocker before closing the linked execution item.",
            )
        )
    for index, signal in enumerate(risk_signals, start=1):
        risks.append(
            ReportRiskRow(
                risk_id=_safe_text(signal.get("risk_id"), default=f"RISK-{index}"),
                category=_safe_text(signal.get("kind"), default="Risk signal").replace("-", " ").title(),
                severity="Medium",
                description=_safe_text(signal.get("summary") or signal.get("kind") or "Execution risk flagged."),
                owner="PMO",
                related_decision_id=_safe_text(signal.get("decision_id"), default=""),
                recommended_follow_up="Review the signal and assign follow-through before the next checkpoint.",
            )
        )
    _add_missing_follow_through_risks(action_rows, decisions, risks)

    workflow_rows: list[ReportWorkflowRow] = []
    transition_map: dict[str, list[str]] = {}
    for item in workflow.transitions:
        source_stage = _safe_text(item.get("from"), default="Unknown")
        target_stage = _safe_text(item.get("to"), default="Unknown")
        transition_map.setdefault(source_stage, []).append(target_stage)
    for index, stage in enumerate(workflow.stages, start=1):
        stage_name = _safe_text(stage.get("name") if isinstance(stage, dict) else stage, default=f"Stage {index}")
        workflow_rows.append(
            ReportWorkflowRow(
                stage=stage_name,
                outbound_transition=_render_list(transition_map.get(stage_name, []), default="No downstream transition recorded"),
                stage_note=_safe_text(workflow.workflow_summary.get("status") or workflow.workflow_summary.get("top_priority_decision") or "Active workflow stage"),
            )
        )

    participation_rows: list[ReportSpeakerRow] = []
    speaker_ratings = result.meeting_scores.speaker_rating if isinstance(result.meeting_scores.speaker_rating, dict) else {}
    for item in result.speaker_dominance.speakers:
        speaker = _safe_text(item.get("speaker"), default="Unknown")
        dominance_ratio = _safe_float(item.get("dominance_ratio"), default=0.0)
        rating = str(speaker_ratings.get(speaker) or speaker_ratings.get("score") or "n/a")
        participation_note = "Primary contributor" if dominance_ratio >= 50 else "Supporting contributor"
        participation_rows.append(
            ReportSpeakerRow(
                speaker=speaker,
                dominance_ratio=dominance_ratio,
                rating=rating,
                participation_note=participation_note,
            )
        )

    visual_rows: list[ReportVisualRow] = [
        ReportVisualRow(
            artifact_id=item.artifact_id,
            artifact_type=_safe_text(item.artifact_type, default="artifact"),
            display_mode=_safe_text(item.display_mode, default="Unknown"),
            time_window=_format_timestamp_range(item.start_time, item.end_time),
            summary=_safe_text(item.content_summary, default="No summary recorded."),
            insight=_safe_text(item.content_insight, default="No insight recorded."),
            confidence=_safe_float(item.confidence),
        )
        for item in result.visual_artifacts
    ]

    trace_rows: list[ReportTraceRow] = [
        ReportTraceRow(
            trace_id=item.trace_id,
            title=_safe_text(item.title, default=item.trace_id),
            owner=_safe_text(item.owner, default="Unassigned"),
            summary=_safe_text(item.summary, default="No summary recorded."),
            rationale=_render_list(item.rationale),
            next_steps=_render_list(item.next_steps),
            related_artifacts=_render_list(item.related_artifacts),
        )
        for item in result.decision_traces
    ]

    cognitive_rating = result.meeting_scores.cognitive_rating or result.attention_sentiment.cognitive_rating or {}
    focus = _coalesce_mapping_value(cognitive_rating, "meeting_focus", "focus")
    clarity = _coalesce_mapping_value(cognitive_rating, "meeting_clarity", "clarity")
    overload_risk = _coalesce_mapping_value(cognitive_rating, "overload_risk", default="n/a")

    report_title = _derive_title_from_text(Path(result.input_video).stem.replace("_", " ").replace("-", " ").title(), "BoardSight Meeting")
    executive_summary = [
        f"BoardSight assessed this meeting as {result.attention_sentiment.overall_sentiment or 'neutral'} with {_format_ratio(result.attention_sentiment.overall_attention)} overall attention coverage.",
        f"{len(decisions)} decisions and {len(action_rows)} execution actions were captured, with execution readiness scored at {result.meeting_scores.execution_readiness:.2f}.",
        f"Top workflow signal: {workflow.workflow_summary.get('top_priority_decision', 'No high-priority decision identified')}.",
        f"Cognitive readout: focus {focus}, clarity {clarity}, overload risk {overload_risk}.",
        _safe_text(result.meeting_scores.meeting_conclusion, default="No executive conclusion was generated."),
    ]

    metadata_rows = [
        ("Input video", _safe_text(result.input_video)),
        ("Analysis profile", _safe_text(metadata.get("analysis_profile"), default="n/a")),
        ("Source mode", _safe_text(metadata.get("source_mode"), default="n/a")),
        ("Runtime profile", _safe_text((metadata.get("performance_report") or {}).get("runtime_profile"), default="n/a")),
        ("Contract version", _safe_text(agentic_contract.get("contract_version") or metadata.get("data_contract_version"), default="n/a")),
        ("Export generated", generated_at),
        ("Speaker count", str(len(participation_rows) or len(result.transcript.speaker_directory))),
        ("Visual evidence windows", str(len(visual_rows))),
        ("Warnings", str(len(result.warnings))),
    ]

    recommended_follow_through: list[str] = []
    if action_rows:
        recommended_follow_through.append(
            f"Confirm owners and due dates for the top {min(3, len(action_rows))} execution items before the next status checkpoint."
        )
    if risks:
        recommended_follow_through.append("Escalate open blockers and risk signals that remain unresolved after this meeting.")
    if visual_rows:
        recommended_follow_through.append("Use the visual evidence appendix to validate presentation-dependent decisions and approvals.")
    if not recommended_follow_through:
        recommended_follow_through.append("No urgent follow-through items were detected; keep the decision register as the audit trail.")

    return StructuredReportModel(
        report_title=report_title,
        generated_at=generated_at,
        executive_summary=executive_summary,
        metadata_rows=metadata_rows,
        decisions=decisions,
        actions=sorted(action_rows, key=lambda item: item.execution_order),
        risks=risks,
        workflow=workflow_rows,
        participation=participation_rows,
        visuals=visual_rows,
        traces=trace_rows,
        recommended_follow_through=recommended_follow_through,
        warnings=list(result.warnings or []),
    )


def _markdown_table(headers: list[str], rows: list[list[Any]]) -> list[str]:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        normalized = [str(item).replace("\n", " ").replace("|", "/") for item in row]
        lines.append("| " + " | ".join(normalized) + " |")
    return lines


def build_markdown_report(result: PipelineResult) -> str:
    model = build_structured_report_model(result)
    lines = [
        "# BoardSight Structured Governance Report",
        "",
        f"**Meeting:** {model.report_title}",
        f"**Export generated:** {model.generated_at}",
        "",
        "## Executive Summary",
    ]
    lines.extend([f"- {item}" for item in model.executive_summary])
    lines.extend(["", "## Meeting Metadata"])
    lines.extend([f"- **{label}:** {value}" for label, value in model.metadata_rows])

    lines.extend(["", "## Decision Register"])
    if model.decisions:
        lines.extend(
            _markdown_table(
                ["Decision ID", "Title", "Owner", "Urgency", "Impact", "Status", "Next Action"],
                [
                    [
                        item.decision_id,
                        item.title,
                        item.owner,
                        item.urgency,
                        item.impact,
                        item.status,
                        item.next_action,
                    ]
                    for item in model.decisions
                ],
            )
        )
    else:
        lines.append("- No decisions were captured.")

    lines.extend(["", "## Action Register"])
    if model.actions:
        lines.extend(
            _markdown_table(
                ["Action ID", "Task", "Owner", "Due Date", "Status", "GitLab Sync"],
                [
                    [
                        item.action_id,
                        item.task_title,
                        item.owner,
                        item.due_date or "n/a",
                        item.status,
                        item.gitlab_sync_status,
                    ]
                    for item in model.actions
                ],
            )
        )
    else:
        lines.append("- No execution actions were captured.")

    lines.extend(["", "## Blockers and Risks"])
    if model.risks:
        lines.extend([f"- **{item.risk_id}** [{item.severity}] {item.category}: {item.description}" for item in model.risks])
    else:
        lines.append("- No blockers or risk signals were identified.")

    lines.extend(["", "## Workflow Stage Summary"])
    if model.workflow:
        lines.extend([f"- **{item.stage}** -> {item.outbound_transition}. {item.stage_note}" for item in model.workflow])
    else:
        lines.append("- No workflow stages were modeled.")

    lines.extend(["", "## Participation and Dominance"])
    if model.participation:
        lines.extend([f"- {item.speaker}: {_format_ratio(item.dominance_ratio)} dominance | {item.participation_note}" for item in model.participation])
    else:
        lines.append("- No speaker participation statistics were available.")

    lines.extend(["", "## Visual Evidence Appendix"])
    if model.visuals:
        lines.extend([f"- {item.artifact_id}: {item.artifact_type} | {item.display_mode} | {item.time_window} | {item.summary}" for item in model.visuals])
    else:
        lines.append("- No visual evidence windows were captured.")

    lines.extend(["", "## Traceability"])
    if model.traces:
        lines.extend([f"- {item.trace_id}: {item.title} | owner {item.owner} | next steps {item.next_steps}" for item in model.traces])
    else:
        lines.append("- No decision traces were available.")

    lines.extend(["", "## Recommended Follow-Through"])
    lines.extend([f"- {item}" for item in model.recommended_follow_through])

    if model.warnings:
        lines.extend(["", "## Warnings"])
        lines.extend([f"- {item}" for item in model.warnings])

    return "\n".join(lines) + "\n"


def _add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = None


def _add_table(document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        _set_cell_text(header_cells[index], header, bold=True)
        _shade_cell(header_cells[index], "E2E8F0")
    for row_values in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_values):
            _set_cell_text(row_cells[index], str(value))


def write_docx_report(result: PipelineResult, path: Path) -> None:
    import docx
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    model = build_structured_report_model(result)
    document = docx.Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10)

    header = section.header.paragraphs[0]
    header.text = "BoardSight | Structured Governance Report"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string("475569")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    _add_page_number(footer)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("BoardSight Structured Governance Report")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0F172A")

    subtitle = document.add_paragraph()
    subtitle.add_run(model.report_title).bold = True
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string("2563EB")
    subtitle.add_run(f"\nExport generated {model.generated_at}")
    subtitle.runs[1].font.size = Pt(9)
    subtitle.runs[1].font.color.rgb = RGBColor.from_string("64748B")

    def add_heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.space_before = Pt(10)
        paragraph.space_after = Pt(6)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string("0F172A")

    add_heading("Executive Summary")
    for item in model.executive_summary:
        document.add_paragraph(item, style="List Bullet")

    add_heading("Meeting Metadata")
    _add_table(document, ["Field", "Value"], [[label, value] for label, value in model.metadata_rows])

    add_heading("Decision Register")
    if model.decisions:
        _add_table(
            document,
            ["Decision ID", "Title", "Owner", "Urgency", "Impact", "Status", "Next Action", "GitLab"],
            [
                [
                    item.decision_id,
                    item.title,
                    item.owner,
                    item.urgency,
                    item.impact,
                    item.status,
                    item.next_action,
                    item.linked_gitlab_issue,
                ]
                for item in model.decisions
            ],
        )
    else:
        document.add_paragraph("No decisions were captured.")

    add_heading("Action Register")
    if model.actions:
        _add_table(
            document,
            ["Action ID", "Task", "Owner", "Due Date", "Status", "Dependencies", "GitLab Sync"],
            [
                [
                    item.action_id,
                    item.task_title,
                    item.owner,
                    item.due_date or "n/a",
                    item.status,
                    item.dependencies,
                    item.gitlab_sync_status,
                ]
                for item in model.actions
            ],
        )
    else:
        document.add_paragraph("No execution actions were captured.")

    add_heading("Blockers and Risks")
    if model.risks:
        _add_table(
            document,
            ["Risk ID", "Category", "Severity", "Description", "Recommended Follow-Up"],
            [
                [
                    item.risk_id,
                    item.category,
                    item.severity,
                    item.description,
                    item.recommended_follow_up,
                ]
                for item in model.risks
            ],
        )
    else:
        document.add_paragraph("No blockers or risk signals were identified.")

    add_heading("Workflow Stage Summary")
    if model.workflow:
        _add_table(
            document,
            ["Stage", "Outbound Transition", "Stage Note"],
            [[item.stage, item.outbound_transition, item.stage_note] for item in model.workflow],
        )
    else:
        document.add_paragraph("No workflow stages were modeled.")

    add_heading("Participation and Dominance Analysis")
    if model.participation:
        _add_table(
            document,
            ["Speaker", "Dominance", "Rating", "Participation Note"],
            [[item.speaker, _format_ratio(item.dominance_ratio), item.rating, item.participation_note] for item in model.participation],
        )
    else:
        document.add_paragraph("No speaker participation statistics were available.")

    add_heading("Visual Evidence Appendix")
    if model.visuals:
        _add_table(
            document,
            ["Artifact", "Type", "Mode", "Time Window", "Summary", "Insight"],
            [[item.artifact_id, item.artifact_type, item.display_mode, item.time_window, item.summary, item.insight] for item in model.visuals],
        )
    else:
        document.add_paragraph("No visual evidence windows were captured.")

    add_heading("Traceability")
    if model.traces:
        _add_table(
            document,
            ["Trace ID", "Title", "Owner", "Summary", "Next Steps"],
            [[item.trace_id, item.title, item.owner, item.summary, item.next_steps] for item in model.traces],
        )
    else:
        document.add_paragraph("No decision traces were available.")

    add_heading("Recommended Follow-Through")
    for item in model.recommended_follow_through:
        document.add_paragraph(item, style="List Bullet")

    if model.warnings:
        add_heading("Warnings")
        for item in model.warnings:
            document.add_paragraph(item, style="List Bullet")

    document.add_section(WD_SECTION.CONTINUOUS)
    document.save(path)


def _pdf_header_footer(canvas, doc) -> None:
    from reportlab.lib import colors

    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor(BRAND_BORDER))
    canvas.setFillColor(colors.HexColor(BRAND_NAVY))
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(doc.leftMargin, doc.pagesize[1] - 30, "BoardSight | Structured Governance Report")
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor(BRAND_MUTED))
    canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, 20, f"Page {doc.page}")
    canvas.restoreState()


def write_pdf_report(result: PipelineResult, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    model = build_structured_report_model(result)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.55 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BoardSightTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor(BRAND_NAVY),
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "BoardSightHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=14,
        textColor=colors.HexColor(BRAND_NAVY),
        spaceBefore=8,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "BoardSightBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor(BRAND_TEXT),
    )
    muted_style = ParagraphStyle(
        "BoardSightMuted",
        parent=body_style,
        textColor=colors.HexColor(BRAND_MUTED),
    )

    def make_table(headers: list[str], rows: list[list[Any]], col_widths: list[float] | None = None) -> Table:
        table_rows = [[Paragraph(f"<b>{header}</b>", body_style) for header in headers]]
        table_rows.extend([[Paragraph(str(item), body_style) for item in row] for row in rows])
        table = Table(table_rows, colWidths=col_widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(BRAND_TEXT)),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(BRAND_BORDER)),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(BRAND_SURFACE)]),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        return table

    story = [
        Paragraph("BoardSight Structured Governance Report", title_style),
        Paragraph(model.report_title, heading_style),
        Paragraph(f"Export generated {model.generated_at}", muted_style),
        Spacer(1, 10),
        Paragraph("Executive Summary", heading_style),
    ]
    for item in model.executive_summary:
        story.append(Paragraph(f"&bull; {item}", body_style))
    story.extend(
        [
            Spacer(1, 8),
            Paragraph("Meeting Metadata", heading_style),
            make_table(["Field", "Value"], [[label, value] for label, value in model.metadata_rows], [120, 360]),
            Spacer(1, 8),
            Paragraph("Decision Register", heading_style),
        ]
    )
    if model.decisions:
        story.append(
            make_table(
                ["Decision ID", "Title", "Owner", "Urgency", "Impact", "Status", "Next Action"],
                [[item.decision_id, item.title, item.owner, item.urgency, item.impact, item.status, item.next_action] for item in model.decisions],
                [58, 140, 62, 48, 45, 60, 122],
            )
        )
    else:
        story.append(Paragraph("No decisions were captured.", body_style))

    story.extend([Spacer(1, 8), Paragraph("Action Register", heading_style)])
    if model.actions:
        story.append(
            make_table(
                ["Action ID", "Task", "Owner", "Due Date", "Status", "GitLab"],
                [[item.action_id, item.task_title, item.owner, item.due_date or "n/a", item.status, item.gitlab_sync_status] for item in model.actions],
                [55, 180, 70, 62, 65, 75],
            )
        )
    else:
        story.append(Paragraph("No execution actions were captured.", body_style))

    story.extend([Spacer(1, 8), Paragraph("Blockers and Risks", heading_style)])
    if model.risks:
        story.append(
            make_table(
                ["Risk ID", "Category", "Severity", "Description", "Follow-Up"],
                [[item.risk_id, item.category, item.severity, item.description, item.recommended_follow_up] for item in model.risks],
                [55, 70, 48, 150, 165],
            )
        )
    else:
        story.append(Paragraph("No blockers or risk signals were identified.", body_style))

    story.extend([Spacer(1, 8), Paragraph("Workflow Stage Summary", heading_style)])
    if model.workflow:
        story.append(
            make_table(
                ["Stage", "Outbound Transition", "Stage Note"],
                [[item.stage, item.outbound_transition, item.stage_note] for item in model.workflow],
                [110, 150, 228],
            )
        )
    else:
        story.append(Paragraph("No workflow stages were modeled.", body_style))

    story.extend([Spacer(1, 8), Paragraph("Participation and Dominance Analysis", heading_style)])
    if model.participation:
        story.append(
            make_table(
                ["Speaker", "Dominance", "Rating", "Participation Note"],
                [[item.speaker, _format_ratio(item.dominance_ratio), item.rating, item.participation_note] for item in model.participation],
                [120, 70, 70, 228],
            )
        )
    else:
        story.append(Paragraph("No speaker participation statistics were available.", body_style))

    story.extend([Spacer(1, 8), Paragraph("Visual Evidence Appendix", heading_style)])
    if model.visuals:
        story.append(
            make_table(
                ["Artifact", "Type", "Mode", "Time", "Summary"],
                [[item.artifact_id, item.artifact_type, item.display_mode, item.time_window, item.summary] for item in model.visuals],
                [48, 54, 70, 70, 246],
            )
        )
    else:
        story.append(Paragraph("No visual evidence windows were captured.", body_style))

    story.extend([Spacer(1, 8), Paragraph("Traceability", heading_style)])
    if model.traces:
        story.append(
            make_table(
                ["Trace ID", "Title", "Owner", "Next Steps"],
                [[item.trace_id, item.title, item.owner, item.next_steps] for item in model.traces],
                [60, 160, 80, 188],
            )
        )
    else:
        story.append(Paragraph("No decision traces were available.", body_style))

    story.extend([Spacer(1, 8), Paragraph("Recommended Follow-Through", heading_style)])
    for item in model.recommended_follow_through:
        story.append(Paragraph(f"&bull; {item}", body_style))
    if model.warnings:
        story.extend([Spacer(1, 8), Paragraph("Warnings", heading_style)])
        for item in model.warnings:
            story.append(Paragraph(f"&bull; {item}", body_style))

    doc.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)


def _auto_fit_excel_columns(worksheet, dataframe) -> None:
    for index, column in enumerate(dataframe.columns):
        series = dataframe[column].astype(str) if not dataframe.empty else []
        max_length = max([len(str(column))] + [len(item) for item in series]) if len(series) else len(str(column))
        worksheet.set_column(index, index, min(max(max_length + 2, 14), 42))


def write_excel_report(result: PipelineResult, path: Path) -> None:
    import pandas as pd

    model = build_structured_report_model(result)

    summary_df = pd.DataFrame(
        [
            {"section": "Executive Summary", "detail": item}
            for item in model.executive_summary
        ]
        + [{"section": label, "detail": value} for label, value in model.metadata_rows]
        + [{"section": "Recommended Follow-Through", "detail": item} for item in model.recommended_follow_through]
    )
    decisions_df = pd.DataFrame([item.__dict__ for item in model.decisions])
    actions_df = pd.DataFrame([item.__dict__ for item in model.actions])
    risks_df = pd.DataFrame([item.__dict__ for item in model.risks])
    workflow_df = pd.DataFrame([item.__dict__ for item in model.workflow])
    speakers_df = pd.DataFrame([item.__dict__ for item in model.participation])
    visuals_df = pd.DataFrame([item.__dict__ for item in model.visuals])
    traces_df = pd.DataFrame([item.__dict__ for item in model.traces])
    transcript_df = pd.DataFrame([segment.__dict__ for segment in result.transcript.segments])

    with pd.ExcelWriter(path, engine="xlsxwriter") as writer:
        workbook = writer.book
        header_format = workbook.add_format(
            {
                "bold": True,
                "font_color": "white",
                "bg_color": BRAND_NAVY,
                "border": 1,
                "valign": "top",
            }
        )
        text_format = workbook.add_format({"text_wrap": True, "valign": "top", "border": 1})

        sheet_map = {
            "summary": summary_df,
            "decisions": decisions_df,
            "actions": actions_df,
            "blockers_risks": risks_df,
            "workflow": workflow_df,
            "speakers": speakers_df,
            "visual_evidence": visuals_df,
            "traceability": traces_df,
            "transcript": transcript_df,
        }

        for sheet_name, dataframe in sheet_map.items():
            dataframe.to_excel(writer, sheet_name=sheet_name, index=False)
            worksheet = writer.sheets[sheet_name]
            worksheet.freeze_panes(1, 0)
            for col_num, value in enumerate(dataframe.columns):
                worksheet.write(0, col_num, value, header_format)
                worksheet.set_column(col_num, col_num, 18, text_format)
            _auto_fit_excel_columns(worksheet, dataframe)


def write_summary_image(result: PipelineResult, path: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    model = build_structured_report_model(result)
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#0F172A")
    ax.axis("off")
    ax.text(0.05, 0.86, "BoardSight Governance Snapshot", fontsize=18, weight="bold", color="white")
    ax.text(0.05, 0.74, model.report_title, fontsize=11, color="#93C5FD")
    ax.text(0.05, 0.58, f"Decisions captured: {len(model.decisions)}", fontsize=12, color="white")
    ax.text(0.05, 0.46, f"Execution actions: {len(model.actions)}", fontsize=12, color="white")
    ax.text(0.05, 0.34, f"Open blockers and risks: {len(model.risks)}", fontsize=12, color="white")
    ax.text(0.05, 0.22, f"Top workflow signal: {_compact_text(result.workflow_model.workflow_summary.get('top_priority_decision', 'None'), 42)}", fontsize=12, color="white")
    ax.text(0.05, 0.1, _compact_text(model.recommended_follow_through[0] if model.recommended_follow_through else "No immediate follow-through.", 88), fontsize=10, color="#CBD5E1")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)


def write_structured_reports(result: PipelineResult, output_dir: Path) -> dict[str, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    files: dict[str, str] = {}

    transcript_csv = output_dir / "transcript.csv"
    with transcript_csv.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["start", "end", "speaker", "text", "confidence"])
        for segment in result.transcript.segments:
            writer.writerow([segment.start, segment.end, segment.speaker, segment.text, segment.confidence])
    files["excel_ready_csv"] = str(transcript_csv)

    report_body = build_markdown_report(result)
    report_md = output_dir / "structured_report.md"
    report_md.write_text(report_body, encoding="utf-8")
    files["markdown_report"] = str(report_md)

    try:
        write_docx_report(result, output_dir / "structured_report.docx")
        files["docx"] = str(output_dir / "structured_report.docx")
    except Exception as exc:
        files["docx_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    try:
        write_pdf_report(result, output_dir / "structured_report.pdf")
        files["pdf"] = str(output_dir / "structured_report.pdf")
    except Exception as exc:
        files["pdf_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    try:
        write_excel_report(result, output_dir / "structured_report.xlsx")
        files["xlsx"] = str(output_dir / "structured_report.xlsx")
    except Exception as exc:
        files["xlsx_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    try:
        write_summary_image(result, output_dir / "summary_card.png")
        files["image"] = str(output_dir / "summary_card.png")
    except Exception as exc:
        files["image_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    return files


def write_structured_report_artifact(result: PipelineResult, output_dir: Path, file_name: str) -> Path | None:
    """Generate only the requested export instead of rebuilding the full report bundle."""
    safe_name = Path(file_name).name
    output_dir.mkdir(parents=True, exist_ok=True)
    artifact_path = output_dir / safe_name

    if safe_name == "structured_report.pdf":
        write_pdf_report(result, artifact_path)
    elif safe_name == "structured_report.docx":
        write_docx_report(result, artifact_path)
    elif safe_name == "structured_report.xlsx":
        write_excel_report(result, artifact_path)
    elif safe_name == "transcript.csv":
        with artifact_path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            writer.writerow(["start", "end", "speaker", "text", "confidence"])
            for segment in result.transcript.segments:
                writer.writerow([segment.start, segment.end, segment.speaker, segment.text, segment.confidence])
    elif safe_name == "summary_card.png":
        write_summary_image(result, artifact_path)
    else:
        return None
    return artifact_path if artifact_path.exists() else None
